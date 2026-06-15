"""
Resume processor for handling PDF, TXT, and Word documents
"""
import os
from PyPDF2 import PdfReader
from docx import Document
import mimetypes

class ResumeProcessor:
    """Handle resume extraction from various file formats"""
    
    ALLOWED_EXTENSIONS = {'pdf', 'txt', 'docx', 'doc'}
    ALLOWED_MIMETYPES = {
        'application/pdf',
        'text/plain',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'application/msword'
    }
    
    @staticmethod
    def allowed_file(filename):
        """Check if file extension is allowed"""
        return '.' in filename and filename.rsplit('.', 1)[1].lower() in ResumeProcessor.ALLOWED_EXTENSIONS
    
    @staticmethod
    def extract_text(filepath, filename):
        """
        Extract text from resume file
        Returns: tuple (success: bool, content: str, error: str)
        """
        try:
            file_ext = filename.rsplit('.', 1)[1].lower()
            
            if file_ext == 'pdf':
                return ResumeProcessor._extract_pdf(filepath)
            elif file_ext == 'txt':
                return ResumeProcessor._extract_txt(filepath)
            elif file_ext in ['docx', 'doc']:
                return ResumeProcessor._extract_docx(filepath)
            else:
                return False, '', f"Unsupported file format: {file_ext}"
                
        except Exception as e:
            return False, '', f"Error processing file: {str(e)}"
    
    @staticmethod
    def _extract_pdf(filepath):
        """Extract text from PDF"""
        try:
            pdf_reader = PdfReader(filepath)
            text = ""
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            if not text.strip():
                return False, '', "PDF is empty or text could not be extracted"
            
            return True, text, None
            
        except Exception as e:
            return False, '', f"Error reading PDF: {str(e)}"
    
    @staticmethod
    def _extract_txt(filepath):
        """Extract text from TXT file"""
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
            
            if not text.strip():
                return False, '', "TXT file is empty"
            
            return True, text, None
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(filepath, 'r', encoding='latin-1') as f:
                    text = f.read()
                return True, text, None
            except Exception as e:
                return False, '', f"Error reading TXT: {str(e)}"
        except Exception as e:
            return False, '', f"Error reading TXT: {str(e)}"
    
    @staticmethod
    def _extract_docx(filepath):
        """Extract text from Word document"""
        try:
            doc = Document(filepath)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += "\n" + cell.text
            
            if not text.strip():
                return False, '', "Word document is empty"
            
            return True, text, None
            
        except Exception as e:
            return False, '', f"Error reading Word document: {str(e)}"
    
    @staticmethod
    def get_candidate_name(filepath, filename):
        """Extract candidate name from filename"""
        name = filename.rsplit('.', 1)[0]
        # Remove common prefixes
        for prefix in ['resume_', 'cv_', 'cv-', 'resume-']:
            if name.lower().startswith(prefix):
                name = name[len(prefix):]
        return name
